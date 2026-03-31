"""
Resume Generator — Exactly 2 pages, ATS-friendly DOCX + PDF.
Specs: 1-inch margins, Calibri, 10.5pt body, 12pt headings, 20pt name,
1.15 line spacing, left-aligned, no tables/icons/images.
"""

import json, re, sys, os, subprocess

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# Colors
NAVY = RGBColor(43, 43, 94)
BLACK = RGBColor(30, 30, 30)
GRAY = RGBColor(100, 100, 100)
DGRAY = RGBColor(60, 60, 60)

# Sizing
BODY = 10.5
HEAD = 12
NAME_SIZE = 20
LINE_SP = 1.15  # multiplier

# A4 with 1-inch margins: usable width = 8.27 - 2 = 6.27 inches
USABLE = 6.27


def load_data(path="resume_data.json"):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def clean(t):
    return re.sub(r" {2,}", " ", re.sub(r"[*_~`#]", "", t)).strip()


def sf(run, size=BODY, bold=False, italic=False, color=BLACK):
    """Set font on a run."""
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def kwn(p):
    p._element.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def kl(p):
    p._element.get_or_add_pPr().append(OxmlElement("w:keepLines"))


def set_line_spacing(p, multiplier=LINE_SP):
    """Set line spacing as a multiplier (1.15 = 115%)."""
    pPr = p._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    # 240 twips = single spacing; multiply for desired ratio
    spacing.set(qn("w:line"), str(int(240 * multiplier)))
    spacing.set(qn("w:lineRule"), "auto")


def add_right_tab(p):
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(USABLE * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


def border_bottom(p, sz="6", color="2B2B5E"):
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color)
    pBdr.append(b)
    pPr.append(pBdr)


def heading(doc, text, sp_before=14, sp_after=4):
    """Section heading: 12pt bold navy, bottom border, keep-with-next."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after = Pt(sp_after)
    set_line_spacing(p)
    r = p.add_run(text.upper())
    sf(r, size=HEAD, bold=True, color=NAVY)
    border_bottom(p)
    kwn(p)
    return p


def body_para(doc, text, sp_before=2, sp_after=2, color=DGRAY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after = Pt(sp_after)
    set_line_spacing(p)
    r = p.add_run(clean(text))
    sf(r, color=color)
    return p


def bullet(doc, text, sp_before=1.5, sp_after=1.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    set_line_spacing(p)
    r = p.add_run("\u2022  " + clean(text))
    sf(r)
    return p


def role_header(doc, company, location, role, period):
    """Two-line role header: Company...Location / Role...Period (right-aligned)."""
    # Line 1: Company + Location
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    set_line_spacing(p)
    kwn(p)
    add_right_tab(p)
    r = p.add_run(clean(company))
    sf(r, bold=True, color=BLACK)
    if location:
        p.add_run("\t")
        r = p.add_run(clean(location))
        sf(r, size=10, italic=True, color=GRAY)

    # Line 2: Role + Period
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    set_line_spacing(p)
    kwn(p)
    add_right_tab(p)
    r = p.add_run(clean(role))
    sf(r, bold=True, color=NAVY)
    p.add_run("\t")
    r = p.add_run(clean(period))
    sf(r, size=10, italic=True, color=GRAY)


def generate_docx(data, out="Harikrishnan_2026.docx"):
    doc = Document()

    # Page setup: A4, 1-inch margins all sides
    for sec in doc.sections:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Default style
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(BODY)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)

    # ===== NAME (centered, 20pt) =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    set_line_spacing(p)
    r = p.add_run(clean(data["name"]).upper())
    sf(r, size=NAME_SIZE, bold=True, color=NAVY)

    # ===== CONTACT (centered) =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_line_spacing(p)
    parts = []
    if data.get("phone"): parts.append(data["phone"])
    if data.get("email"): parts.append(data["email"])
    if data.get("linkedin"): parts.append(data["linkedin"])
    r = p.add_run("   |   ".join(parts))
    sf(r, size=9.5, color=GRAY)

    # Thick divider
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    border_bottom(p, sz="10")

    # ===== PROFESSIONAL SUMMARY =====
    heading(doc, "Professional Summary")
    body_para(doc, data["summary"], sp_before=3, sp_after=4)

    # ===== SKILLS =====
    heading(doc, "Skills")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    set_line_spacing(p)
    r = p.add_run("  |  ".join(clean(sk) for sk in data["skills"]))
    sf(r, size=10, color=DGRAY)

    # ===== WORK EXPERIENCE =====
    heading(doc, "Work Experience")

    for job in data["experience"]:
        role_header(doc, job["company"], job.get("location", ""),
                    job["role"], job["period"])

        if job.get("categories"):
            for cat in job["categories"]:
                # Category sub-heading
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                set_line_spacing(p)
                kwn(p)
                r = p.add_run(clean(cat["name"]))
                sf(r, size=10, bold=True, italic=True, color=DGRAY)
                for b in cat["bullets"]:
                    bullet(doc, b)
        elif job.get("bullets"):
            for b in job["bullets"]:
                bullet(doc, b)

    # ===== KEY PROJECTS =====
    if data.get("projects"):
        heading(doc, "Key Projects")
        for i, proj in enumerate(data["projects"], 1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            set_line_spacing(p)
            kwn(p)
            r = p.add_run(f"{i}. {clean(proj['title'])}")
            sf(r, bold=True, color=BLACK)

            for key in ("problem", "action", "impact"):
                if proj.get(key):
                    bp = doc.add_paragraph()
                    bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    bp.paragraph_format.space_before = Pt(1)
                    bp.paragraph_format.space_after = Pt(1)
                    bp.paragraph_format.left_indent = Inches(0.25)
                    set_line_spacing(bp)
                    lbl = bp.add_run(key.capitalize() + ": ")
                    sf(lbl, size=10, bold=True, color=NAVY)
                    txt = bp.add_run(clean(proj[key]))
                    sf(txt, size=10, color=DGRAY)

    # ===== EDUCATION =====
    heading(doc, "Education")
    for edu in data["education"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        set_line_spacing(p)
        kl(p)
        add_right_tab(p)

        r = p.add_run(clean(edu["degree"]))
        sf(r, bold=True, color=BLACK)
        r = p.add_run(" \u2014 ")
        sf(r, color=GRAY)
        r = p.add_run(clean(edu["institution"]))
        sf(r, color=DGRAY)
        if edu.get("grade"):
            r = p.add_run(f" ({clean(edu['grade'])})")
            sf(r, size=9.5, color=GRAY)
        p.add_run("\t")
        r = p.add_run(clean(edu["period"]))
        sf(r, size=10, italic=True, color=GRAY)

    doc.save(out)
    return out


def generate_pdf(docx_path, pdf_path="Harikrishnan_2026.pdf"):
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        wdoc = word.Documents.Open(os.path.abspath(docx_path))
        wdoc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        wdoc.Close()
        word.Quit()
        return pdf_path
    except Exception:
        print("PDF: convert DOCX manually via Word.")
        return None


def count_pages(docx_path):
    """Use Word COM to count pages in the generated doc."""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        wdoc = word.Documents.Open(os.path.abspath(docx_path))
        wdoc.Repaginate()
        pages = wdoc.ComputeStatistics(2)  # wdStatisticPages = 2
        wdoc.Close(False)
        word.Quit()
        return pages
    except Exception:
        return None


def main():
    data = load_data()
    docx_path = generate_docx(data)
    pages = count_pages(docx_path)
    if pages is not None:
        print(f"Page count: {pages}")
        if pages != 2:
            print(f"WARNING: Document is {pages} page(s), target is 2.")
    generate_pdf(docx_path)
    print("Done.")


if __name__ == "__main__":
    main()
